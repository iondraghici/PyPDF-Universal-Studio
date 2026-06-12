#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PyPDF Studio V2 - Procesor Universal de Documente
Fix complet: incarcare DOCX, preview text, conversie DOCX->PDF via docx2pdf
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import sys
import re
import subprocess
from datetime import datetime

# ------------------------------------------------------------------
# IMPORTURI
# ------------------------------------------------------------------
try:
    from pypdf import PdfReader, PdfWriter
    PYPDF_OK = True
except ImportError:
    PYPDF_OK = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from PIL import Image, ImageDraw, ImageTk
    PIL_OK = True
except ImportError:
    PIL_OK = False

try:
    from pdf2docx import Converter
    PDF2DOCX_OK = True
except ImportError:
    PDF2DOCX_OK = False

try:
    from docx2pdf import convert
    DOCX2PDF_OK = True
except ImportError:
    DOCX2PDF_OK = False

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_OK = True
except ImportError:
    OCR_OK = False


# ------------------------------------------------------------------
# CLASA PRINCIPALĂ
# ------------------------------------------------------------------
class PyPDFStudioV2(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("PyPDF Studio V2 - Procesor Universal de Documente")
        self.geometry("1280x900")
        self.minsize(1100, 800)

        # Stare
        self.current_file = None
        self.current_doc_type = None
        self.signature_img = None
        self.manual_fields = []
        self.field_counter = 0
        self.pdf_doc = None
        self.current_page = 0
        self.zoom = 1.5
        self.sig_x = self.sig_y = self.sig_page = None
        self.click_mode = tk.StringVar(value="signature")
        self.pan_start_x = 0
        self.pan_start_y = 0

        self.status_var = tk.StringVar(value="Gata")
        self._build_styles()
        self._build_ui()
        self._check_deps()

    def _build_styles(self):
        self.style = ttk.Style(self)
        self.style.theme_use('clam')
        self.style.configure('TButton', font=('Segoe UI', 10))
        self.style.configure('TLabel', font=('Segoe UI', 10))
        self.style.configure('Header.TLabel', font=('Segoe UI', 13, 'bold'))
        self.style.configure('Action.TButton', font=('Segoe UI', 10, 'bold'))

    def _check_deps(self):
        lipsa = []
        if not PYPDF_OK: lipsa.append("pypdf")
        if not PYMUPDF_OK: lipsa.append("PyMuPDF")
        if not DOCX_OK: lipsa.append("python-docx")
        if not PIL_OK: lipsa.append("Pillow")
        if lipsa:
            self.status_var.set(f"CRITIC: {', '.join(lipsa)} lipsesc!")
            messagebox.showwarning("Dependente lipsa", f"Biblioteci esentiale lipsesc: {', '.join(lipsa)}.")

    def _build_ui(self):
        toolbar = ttk.Frame(self)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=8, pady=6)
        ttk.Label(toolbar, text="PyPDF Studio V2", font=('Segoe UI', 16, 'bold')).pack(side=tk.LEFT, padx=5)
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        ttk.Button(toolbar, text="Despre", command=self._about).pack(side=tk.RIGHT, padx=5)

        ttk.Label(self, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W, font=('Segoe UI', 9)).pack(side=tk.BOTTOM, fill=tk.X)

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(expand=True, fill="both", padx=10, pady=5)

        self._tab_merge()
        self._tab_split()
        self._tab_fill_sign()
        self._tab_security()
        self._tab_convert()

    # ================================================================
    # TAB 1: MERGE
    # ================================================================
    def _tab_merge(self):
        self.tab_merge = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_merge, text="  Unire PDF  ")
        f = ttk.Frame(self.tab_merge, padding=15)
        f.pack(expand=True, fill="both")
        ttk.Label(f, text="Unire fișiere PDF", style='Header.TLabel').pack(anchor=tk.W)
        ttk.Label(f, text="Selectați fișierele și ordonati-le cu sagetile.").pack(anchor=tk.W, pady=(0,10))

        lst_frm = ttk.Frame(f)
        lst_frm.pack(fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(lst_frm)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.merge_lb = tk.Listbox(lst_frm, yscrollcommand=sb.set, font=('Consolas', 10), selectmode=tk.EXTENDED)
        self.merge_lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.config(command=self.merge_lb.yview)

        self.merge_files = []
        bf = ttk.Frame(f)
        bf.pack(fill=tk.X, pady=8)
        ttk.Button(bf, text="➕ Adaugă", command=self._merge_add).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="⬆️ Sus", command=self._merge_up).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="⬇️ Jos", command=self._merge_down).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="🗑️ Elimină", command=self._merge_del).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="💾 EXECUTĂ UNIRE", command=self._merge_run, style='Action.TButton').pack(side=tk.RIGHT, padx=2)

    def _merge_add(self):
        files = filedialog.askopenfilenames(title="Selectați PDF-uri", filetypes=[("PDF", "*.pdf")])
        for p in files:
            if p not in self.merge_files:
                self.merge_files.append(p)
                self.merge_lb.insert(tk.END, os.path.basename(p))
        self.status_var.set(f"{len(self.merge_files)} fișiere")

    def _merge_del(self):
        for i in reversed(self.merge_lb.curselection()):
            self.merge_lb.delete(i); self.merge_files.pop(i)

    def _merge_up(self):
        sel = self.merge_lb.curselection()
        if not sel or sel[0] == 0: return
        i = sel[0]
        self.merge_files[i], self.merge_files[i-1] = self.merge_files[i-1], self.merge_files[i]
        self._merge_refresh(); self.merge_lb.selection_set(i-1)

    def _merge_down(self):
        sel = self.merge_lb.curselection()
        if not sel or sel[0] >= len(self.merge_files)-1: return
        i = sel[0]
        self.merge_files[i], self.merge_files[i+1] = self.merge_files[i+1], self.merge_files[i]
        self._merge_refresh(); self.merge_lb.selection_set(i+1)

    def _merge_refresh(self):
        self.merge_lb.delete(0, tk.END)
        for p in self.merge_files: self.merge_lb.insert(tk.END, os.path.basename(p))

    def _merge_run(self):
        if len(self.merge_files) < 2:
            return messagebox.showwarning("Atenție", "Minim 2 fișiere!")
        out = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not out: return
        try:
            w = PdfWriter()
            for p in self.merge_files: w.append(p)
            with open(out, 'wb') as f: w.write(f)
            self.status_var.set("Unire reusita")
            messagebox.showinfo("Succes", f"Salvat:\n{out}")
        except Exception as e:
            messagebox.showerror("Eroare", str(e))

    # ================================================================
    # TAB 2: SPLIT
    # ================================================================
    def _tab_split(self):
        self.tab_split = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_split, text="  Divizare PDF  ")
        f = ttk.Frame(self.tab_split, padding=15)
        f.pack(expand=True, fill="both")
        ttk.Label(f, text="Divizare fișier PDF", style='Header.TLabel').pack(anchor=tk.W)
        ff = ttk.Frame(f); ff.pack(fill=tk.X, pady=5)
        self.split_path_var = tk.StringVar()
        ttk.Entry(ff, textvariable=self.split_path_var, state='readonly').pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(ff, text="📂 Încarcă PDF", command=self._split_load).pack(side=tk.LEFT, padx=5)
        self.split_info = ttk.Label(f, text="Niciun fișier", foreground="blue")
        self.split_info.pack(anchor=tk.W, pady=5)
        opts = ttk.LabelFrame(f, text="Opțiuni", padding=10)
        opts.pack(fill=tk.X, pady=10)
        self.split_mode = tk.StringVar(value="range")
        ttk.Radiobutton(opts, text="Intervale (ex: 1-3,5,7-9)", variable=self.split_mode, value="range").pack(anchor=tk.W)
        ttk.Radiobutton(opts, text="Fiecare pagină separat", variable=self.split_mode, value="single").pack(anchor=tk.W)
        ttk.Radiobutton(opts, text="La fiecare N pagini", variable=self.split_mode, value="every").pack(anchor=tk.W)
        self.split_entry = ttk.Entry(opts, width=50)
        self.split_entry.pack(anchor=tk.W, padx=20, pady=5)
        ttk.Button(f, text="✂️ EXECUTĂ DIVIZAREA", command=self._split_run, style='Action.TButton').pack(pady=15)

    def _split_load(self):
        p = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if not p: return
        self.split_file = p; self.split_path_var.set(p)
        try:
            r = PdfReader(p); self.split_total = len(r.pages)
            self.split_info.config(text=f"Pagini: {self.split_total}")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    def _split_run(self):
        if not getattr(self, 'split_file', None):
            return messagebox.showwarning("Atenție", "Încărcați un PDF!")
        d = filedialog.askdirectory(title="Folder salvare")
        if not d: return
        try:
            r = PdfReader(self.split_file); n = len(r.pages)
            base = os.path.splitext(os.path.basename(self.split_file))[0]
            mode = self.split_mode.get()
            if mode == "single":
                for i in range(n):
                    w = PdfWriter(); w.add_page(r.pages[i])
                    with open(os.path.join(d, f"{base}_p{i+1}.pdf"), 'wb') as f: w.write(f)
                msg = f"{n} fișiere create"
            elif mode == "every":
                val = self.split_entry.get().strip()
                if not val.isdigit(): return messagebox.showerror("Eroare", "N invalid!")
                step = int(val); idx = 1
                for s in range(0, n, step):
                    w = PdfWriter()
                    for i in range(s, min(s+step, n)): w.add_page(r.pages[i])
                    with open(os.path.join(d, f"{base}_part{idx}.pdf"), 'wb') as f: w.write(f)
                    idx += 1
                msg = f"{idx-1} părți"
            else:
                parts = [x.strip() for x in self.split_entry.get().split(",")]
                idx = 1
                for part in parts:
                    w = PdfWriter()
                    if '-' in part:
                        a, b = part.split('-'); a, b = int(a)-1, int(b)-1
                    else: a = b = int(part)-1
                    for i in range(a, b+1):
                        if 0 <= i < n: w.add_page(r.pages[i])
                    with open(os.path.join(d, f"{base}_seg{idx}.pdf"), 'wb') as f: w.write(f)
                    idx += 1
                msg = "Divizare finalizată"
            self.status_var.set(msg); messagebox.showinfo("Succes", msg)
        except Exception as e: messagebox.showerror("Eroare", str(e))

    # ================================================================
    # TAB 3: FILL & SIGN
    # ================================================================
    def _tab_fill_sign(self):
        self.tab_fill = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_fill, text="  Completare & Semnătură  ")

        # Panou stânga - grid
        left = ttk.Frame(self.tab_fill, padding=10, width=480)
        left.pack(side=tk.LEFT, fill=tk.Y)
        left.pack_propagate(False)
        left.grid_rowconfigure(4, weight=1)
        left.grid_columnconfigure(0, weight=1)

        ttk.Label(left, text="Completare Documente", style='Header.TLabel').grid(row=0, column=0, sticky="w", pady=(0,5))

        lf = ttk.Frame(left)
        lf.grid(row=1, column=0, sticky="ew", pady=3)
        lf.grid_columnconfigure(0, weight=1)
        self.fill_path_var = tk.StringVar()
        ttk.Entry(lf, textvariable=self.fill_path_var, state='readonly').grid(row=0, column=0, sticky="ew")
        ttk.Button(lf, text="📂 PDF/DOCX", command=self._fill_load).grid(row=0, column=1, padx=(5,0))

        mf = ttk.LabelFrame(left, text="Mod click pe preview", padding=8)
        mf.grid(row=2, column=0, sticky="ew", pady=8)
        mf.grid_columnconfigure(0, weight=1)
        ttk.Label(mf, text="Selectați ce plasați la click:").grid(row=0, column=0, sticky="w")
        self.click_combo = ttk.Combobox(mf, textvariable=self.click_mode, state="readonly", width=35)
        self.click_combo.grid(row=1, column=0, sticky="ew", pady=3)
        self.click_combo.set("Semnătură")
        ttk.Label(mf, text="💡 Apasați 📍 pe un câmp, apoi click pe preview",
                 foreground="blue", wraplength=420).grid(row=2, column=0, sticky="w")

        bf = ttk.Frame(left)
        bf.grid(row=3, column=0, sticky="ew", pady=5)
        ttk.Button(bf, text="➕ Adaugă câmp", command=self._field_add_manual).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="🔍 Scanează auto", command=self._field_scan).pack(side=tk.LEFT, padx=2)
        ttk.Button(bf, text="🗑️ Șterge toate", command=self._field_clear).pack(side=tk.LEFT, padx=2)

        cf = ttk.LabelFrame(left, text="Câmpuri de completat", padding=5)
        cf.grid(row=4, column=0, sticky="nsew", pady=5)
        cf.grid_rowconfigure(0, weight=1)
        cf.grid_columnconfigure(0, weight=1)

        cvs = tk.Canvas(cf, highlightthickness=0)
        sb = ttk.Scrollbar(cf, orient=tk.VERTICAL, command=cvs.yview)
        self.fields_frame = ttk.Frame(cvs)
        self.fields_frame.bind("<Configure>", lambda e: cvs.configure(scrollregion=cvs.bbox("all")))
        cvs.create_window((0,0), window=self.fields_frame, anchor="nw")
        cvs.configure(yscrollcommand=sb.set)
        cvs.grid(row=0, column=0, sticky="nsew")
        sb.grid(row=0, column=1, sticky="ns")

        sf = ttk.LabelFrame(left, text="Semnătură", padding=8)
        sf.grid(row=5, column=0, sticky="ew", pady=5)
        self.sig_lbl = ttk.Label(sf, text="Semnătură: Neconfigurată", foreground="red")
        self.sig_lbl.pack(anchor=tk.W)
        ttk.Button(sf, text="✏️ Desenează", command=self._sig_draw).pack(side=tk.LEFT, padx=2)
        ttk.Button(sf, text="📎 Încarcă imagine", command=self._sig_load_img).pack(side=tk.LEFT, padx=2)

        ttk.Button(left, text="💾 Salvează document completat", command=self._fill_save,
                  style='Action.TButton').grid(row=6, column=0, sticky="ew", pady=10)

        # Panou dreapta
        right = ttk.Frame(self.tab_fill, padding=10)
        right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)

        ttk.Label(right, text="Previzualizare (scroll=zoom | drag-dreapta=muta | click-stanga=plaseaza)",
                 style='Header.TLabel').grid(row=0, column=0, sticky="w")

        prev_container = ttk.Frame(right)
        prev_container.grid(row=1, column=0, sticky="nsew", pady=(5,0))
        prev_container.grid_rowconfigure(0, weight=1)
        prev_container.grid_columnconfigure(0, weight=1)

        self.prev_canvas = tk.Canvas(prev_container, bg="#e0e0e0", relief=tk.SUNKEN, bd=2,
                                     xscrollincrement=1, yscrollincrement=1)
        self.hbar = ttk.Scrollbar(prev_container, orient=tk.HORIZONTAL, command=self.prev_canvas.xview)
        self.vbar = ttk.Scrollbar(prev_container, orient=tk.VERTICAL, command=self.prev_canvas.yview)
        self.prev_canvas.configure(xscrollcommand=self.hbar.set, yscrollcommand=self.vbar.set)
        self.hbar.grid(row=1, column=0, sticky="ew")
        self.vbar.grid(row=0, column=1, sticky="ns")
        self.prev_canvas.grid(row=0, column=0, sticky="nsew")

        self.prev_canvas.bind("<MouseWheel>", self._preview_zoom_wheel)
        self.prev_canvas.bind("<Button-4>", self._preview_zoom_wheel)
        self.prev_canvas.bind("<Button-5>", self._preview_zoom_wheel)
        self.prev_canvas.bind("<ButtonPress-3>", self._preview_pan_start)
        self.prev_canvas.bind("<B3-Motion>", self._preview_pan_move)
        self.prev_canvas.bind("<Button-1>", self._preview_click)

        nav = ttk.Frame(right)
        nav.grid(row=2, column=0, sticky="ew", pady=3)
        ttk.Button(nav, text="◀ Pag. ant.", command=lambda: self._preview_nav(-1)).pack(side=tk.LEFT)
        self.page_lbl = ttk.Label(nav, text="Pagina: - / -")
        self.page_lbl.pack(side=tk.LEFT, padx=10)
        ttk.Button(nav, text="Pag. urm. ▶", command=lambda: self._preview_nav(1)).pack(side=tk.LEFT)
        ttk.Separator(nav, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        ttk.Button(nav, text="− Zoom", command=self._zoom_out).pack(side=tk.LEFT)
        self.zoom_lbl = ttk.Label(nav, text="150%", width=6)
        self.zoom_lbl.pack(side=tk.LEFT, padx=3)
        ttk.Button(nav, text="Zoom +", command=self._zoom_in).pack(side=tk.LEFT)
        ttk.Separator(nav, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        ttk.Label(nav, text="Zoom:").pack(side=tk.LEFT)
        self.zoom_var = tk.StringVar(value="1.5")
        z = ttk.Combobox(nav, textvariable=self.zoom_var,
                        values=["0.5","0.8","1.0","1.2","1.5","2.0","2.5","3.0","4.0"],
                        width=5, state="readonly")
        z.pack(side=tk.LEFT, padx=3)
        z.bind("<<ComboboxSelected>>", lambda e: self._preview_render())
        ttk.Button(nav, text="↺ Reset", command=self._zoom_reset).pack(side=tk.LEFT, padx=5)

    def _fill_load(self):
        p = filedialog.askopenfilename(filetypes=[("Documente", "*.pdf *.docx"), ("PDF", "*.pdf"), ("DOCX", "*.docx")])
        if not p: return

        self.current_file = p
        self.current_doc_type = "pdf" if p.lower().endswith(".pdf") else "docx"
        self.fill_path_var.set(p)
        self.current_page = 0
        self.sig_x = self.sig_y = self.sig_page = None
        self._field_clear()

        # Curăță documentul anterior
        if self.pdf_doc:
            try:
                self.pdf_doc.close()
            except Exception:
                pass
            self.pdf_doc = None

        # Încarcă documentul
        try:
            if self.current_doc_type == "pdf" and PYMUPDF_OK:
                self.pdf_doc = fitz.open(p)
                self.status_var.set(f"PDF încărcat: {len(self.pdf_doc)} pagini")
            elif self.current_doc_type == "docx":
                self._load_docx_preview(p)
            else:
                self.status_var.set("Format necunoscut")
        except Exception as e:
            messagebox.showerror("Eroare la încărcare", str(e))
            self.status_var.set(f"Eroare: {str(e)[:50]}")

        self._update_click_combo()
        self._preview_render()

    def _load_docx_preview(self, path):
        """Încarcă DOCX pentru preview - NU folosește PyMuPDF (nu suportă DOCX)"""
        self.docx_preview_text = []

        if DOCX_OK:
            try:
                doc = Document(path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        self.docx_preview_text.append(para.text)
                self.status_var.set(f"DOCX încărcat: {len(self.docx_preview_text)} paragrafe")
            except Exception as e:
                self.docx_preview_text = [f"Eroare citire DOCX: {str(e)}"]
                self.status_var.set("Eroare DOCX")
        else:
            self.docx_preview_text = ["python-docx nu este instalat.", "Rulați: pip install python-docx"]
            self.status_var.set("python-docx lipsă")

    def _update_click_combo(self):
        items = ["Semnătură"]
        for f in self.manual_fields:
            lbl = f["label"].get().strip() or f"Câmp {f['id'].split('_')[1]}"
            items.append(f"{f['id']}|{lbl}")
        self.click_combo['values'] = items
        self.click_mode.set("Semnătură")

    # ------------------------------------------------------------------
    # ZOOM & PANNING
    # ------------------------------------------------------------------
    def _zoom_in(self):
        self.zoom = min(self.zoom + 0.25, 5.0)
        self.zoom_var.set(f"{self.zoom:.1f}")
        self._preview_render()

    def _zoom_out(self):
        self.zoom = max(self.zoom - 0.25, 0.25)
        self.zoom_var.set(f"{self.zoom:.1f}")
        self._preview_render()

    def _zoom_reset(self):
        self.zoom = 1.5
        self.zoom_var.set("1.5")
        self._preview_render()

    def _preview_zoom_wheel(self, event):
        if event.num == 4 or event.delta > 0:
            self.zoom = min(self.zoom + 0.15, 5.0)
        elif event.num == 5 or event.delta < 0:
            self.zoom = max(self.zoom - 0.15, 0.25)
        self.zoom_var.set(f"{self.zoom:.1f}")
        self._preview_render()
        return "break"

    def _preview_pan_start(self, event):
        self.prev_canvas.config(cursor="fleur")
        self.pan_start_x = event.x
        self.pan_start_y = event.y

    def _preview_pan_move(self, event):
        dx = self.pan_start_x - event.x
        dy = self.pan_start_y - event.y
        self.prev_canvas.xview_scroll(dx, "units")
        self.prev_canvas.yview_scroll(dy, "units")
        self.pan_start_x = event.x
        self.pan_start_y = event.y

    # ------------------------------------------------------------------
    # RANDARE PREVIEW
    # ------------------------------------------------------------------
    def _preview_render(self):
        self.prev_canvas.delete("all")

        # === CAZ 1: PDF cu PyMuPDF ===
        if self.pdf_doc and PYMUPDF_OK and self.current_doc_type == "pdf":
            try:
                self.zoom = float(self.zoom_var.get())
                self.zoom_lbl.config(text=f"{int(self.zoom*100)}%")
                page = self.pdf_doc[self.current_page]
                mat = fitz.Matrix(self.zoom, self.zoom)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                self.preview_imgtk = ImageTk.PhotoImage(img)

                self.prev_canvas.config(scrollregion=(0, 0, pix.width, pix.height))
                self.prev_canvas.create_image(0, 0, anchor=tk.NW, image=self.preview_imgtk, tags="page")

                # Câmpuri plasate
                for f in self.manual_fields:
                    if f.get("placed") and f.get("page", 0) == self.current_page and f.get("x") is not None:
                        x, y = f["x"], f["y"]
                        val = f["value"].get().strip()
                        txt = val if val else f["label"].get().strip() or "?"
                        self.prev_canvas.create_rectangle(x-4, y-4, x+4, y+4, fill="green", outline="black")
                        self.prev_canvas.create_text(x, y-14, text=txt[:25], fill="green",
                                                    font=('Arial', 9, 'bold'), anchor=tk.S)

                # Semnătură
                if self.sig_x is not None and self.sig_y is not None and self.sig_page == self.current_page:
                    self.prev_canvas.create_oval(self.sig_x-8, self.sig_y-8, self.sig_x+8, self.sig_y+8,
                                                fill="red", outline="yellow", width=2)
                    self.prev_canvas.create_text(self.sig_x, self.sig_y-18, text="SEMNA",
                                                fill="red", font=('Arial', 11, 'bold'))

                self.page_lbl.config(text=f"Pagina: {self.current_page+1} / {len(self.pdf_doc)}")
                return
            except Exception as e:
                self.prev_canvas.create_text(200, 200, text=f"Eroare preview PDF:\n{str(e)}", fill="red")
                return

        # === CAZ 2: DOCX - afișare text ===
        if self.current_doc_type == "docx" and hasattr(self, 'docx_preview_text'):
            self.prev_canvas.config(scrollregion=(0, 0, 900, max(600, len(self.docx_preview_text)*18 + 80)))
            self.prev_canvas.create_text(20, 20, text="[PREVIZUALIZARE DOCX - MOD TEXT]",
                                        fill="blue", font=('Arial', 12, 'bold'), anchor=tk.NW)
            self.prev_canvas.create_text(20, 45, text="Notă: Pentru preview vizual, convertiți DOCX în PDF mai întâi.",
                                        fill="gray", font=('Arial', 9), anchor=tk.NW)
            y = 70
            for line in self.docx_preview_text[:100]:
                self.prev_canvas.create_text(20, y, text=line[:130], fill="black",
                                            font=('Consolas', 9), anchor=tk.NW)
                y += 16
            self.page_lbl.config(text="Pagina: 1 / 1 (DOCX)")
            return

        # === CAZ 3: Nimic ===
        self.prev_canvas.config(scrollregion=(0, 0, 400, 300))
        self.prev_canvas.create_text(200, 150, text="Previzualizare indisponibilă\nÎncărcați un document",
                                    fill="red", font=('Arial', 12))
        self.page_lbl.config(text="Pagina: - / -")

    def _preview_nav(self, delta):
        if not self.pdf_doc: return
        new = self.current_page + delta
        if 0 <= new < len(self.pdf_doc):
            self.current_page = new
            self._preview_render()

    def _preview_click(self, event):
        if not self.current_file:
            return
        x = self.prev_canvas.canvasx(event.x)
        y = self.prev_canvas.canvasy(event.y)
        mode = self.click_mode.get()

        if mode == "Semnătură":
            self.sig_x = x
            self.sig_y = y
            self.sig_page = self.current_page if self.pdf_doc else 0
            self.status_var.set(f"Semnătura plasată la ({int(x/self.zoom)}, {int(y/self.zoom)}), pag. {self.current_page+1}")
        else:
            fid = mode.split("|")[0]
            for f in self.manual_fields:
                if f["id"] == fid:
                    f["x"] = x
                    f["y"] = y
                    f["page"] = self.current_page if self.pdf_doc else 0
                    f["placed"] = True
                    lbl = f["label"].get().strip() or f"Câmp {fid.split('_')[1]}"
                    self.status_var.set(f"'{lbl}' plasat la ({int(x/self.zoom)}, {int(y/self.zoom)}), pag. {self.current_page+1}")
                    break
        self._preview_render()

    # ------------------------------------------------------------------
    # CÂMPURI MANUALE
    # ------------------------------------------------------------------
    def _field_add_manual(self):
        self.field_counter += 1
        fid = f"fld_{self.field_counter}"
        frame = ttk.Frame(self.fields_frame)
        frame.pack(fill=tk.X, pady=2)

        ttk.Label(frame, text=f"Câmp {self.field_counter}:", width=10).pack(side=tk.LEFT)
        lbl = ttk.Entry(frame, width=16)
        lbl.pack(side=tk.LEFT, padx=2)
        lbl.insert(0, f"Eticheta_{self.field_counter}")

        val = ttk.Entry(frame, width=20)
        val.pack(side=tk.LEFT, padx=2)

        def set_mode():
            self.click_mode.set(f"{fid}|{lbl.get().strip() or f'Câmp {self.field_counter}'}")
            self.click_combo.set(self.click_mode.get())
            self.status_var.set(f"Mod: Plasează '{lbl.get().strip() or f'Câmp {self.field_counter}'}' - click pe preview")

        ttk.Button(frame, text="📍", width=3, command=set_mode).pack(side=tk.LEFT, padx=1)
        ttk.Button(frame, text="✕", width=3, command=lambda: self._field_remove(frame, fid)).pack(side=tk.LEFT, padx=1)

        self.manual_fields.append({
            "id": fid, "frame": frame, "label": lbl, "value": val,
            "x": None, "y": None, "page": 0, "placed": False
        })
        self._update_click_combo()

    def _field_remove(self, frame, fid):
        frame.destroy()
        self.manual_fields = [x for x in self.manual_fields if x["id"] != fid]
        self._update_click_combo()

    def _field_clear(self):
        for f in self.manual_fields: f["frame"].destroy()
        self.manual_fields.clear()
        self.field_counter = 0
        self.sig_x = self.sig_y = self.sig_page = None
        self._update_click_combo()

    def _field_scan(self):
        if not self.current_file:
            return messagebox.showwarning("Atenție", "Încărcați un document!")
        self._field_clear()
        found = []
        try:
            if self.current_doc_type == "pdf":
                r = PdfReader(self.current_file)
                text = ""
                for p in r.pages: text += (p.extract_text() or "") + "\n"
                found = self._extract_patterns(text)
            elif self.current_doc_type == "docx" and DOCX_OK:
                doc = Document(self.current_file)
                text = ""
                for p in doc.paragraphs: text += p.text + "\n"
                found = self._extract_patterns(text)
        except Exception as e:
            messagebox.showerror("Eroare scanare", str(e))
            return

        if not found:
            return messagebox.showinfo("Rezultat", "Nu s-au detectat câmpuri. Folosiți 'Adaugă câmp'.")

        for label, _ in found:
            self.field_counter += 1
            fid = f"fld_{self.field_counter}"
            frame = ttk.Frame(self.fields_frame)
            frame.pack(fill=tk.X, pady=2)
            ttk.Label(frame, text=f"Câmp {self.field_counter}:", width=10).pack(side=tk.LEFT)
            lbl = ttk.Entry(frame, width=16)
            lbl.pack(side=tk.LEFT, padx=2)
            lbl.insert(0, label)
            val = ttk.Entry(frame, width=20)
            val.pack(side=tk.LEFT, padx=2)
            def make_setter(fid=fid, lbl=lbl, fc=self.field_counter):
                return lambda: self._set_click_mode(fid, lbl.get().strip() or f"Câmp {fc}")
            ttk.Button(frame, text="📍", width=3, command=make_setter()).pack(side=tk.LEFT, padx=1)
            ttk.Button(frame, text="✕", width=3, command=lambda f=frame, i=fid: self._field_remove(f, i)).pack(side=tk.LEFT, padx=1)
            self.manual_fields.append({"id": fid, "frame": frame, "label": lbl, "value": val,
                                      "x": None, "y": None, "page": 0, "placed": False})
        self._update_click_combo()
        self.status_var.set(f"{len(found)} câmpuri detectate")

    def _set_click_mode(self, fid, label_text):
        self.click_mode.set(f"{fid}|{label_text}")
        self.click_combo.set(self.click_mode.get())
        self.status_var.set(f"Mod: Plasează '{label_text}' - click pe preview")

    def _extract_patterns(self, text):
        results = []
        patterns = [
            r'(?i)(nume\s*(?:si|și|&)?\s*prenume[\s:\.]*)',
            r'(?i)(cnp[\s:\.]*)',
            r'(?i)(cod\s*numeric\s*personal[\s:\.]*)',
            r'(?i)(serie\s*(?:ci|bi|pasaport)[\s:\.]*)',
            r'(?i)(nr\.?\s*(?:ci|bi)[\s:\.]*)',
            r'(?i)(adresa[\s:\.]*)',
            r'(?i)(domiciliu[\s:\.]*)',
            r'(?i)(localitate[\s:\.]*)',
            r'(?i)(judet[\s:\.]*)',
            r'(?i)(sector[\s:\.]*)',
            r'(?i)(telefon[\s:\.]*)',
            r'(?i)(e-?mail[\s:\.]*)',
            r'(?i)(data[\s:\.]*)',
            r'(?i)(semnat[ura]*[\s:\.]*)',
            r'(?i)(functia[\s:\.]*)',
            r'(?i)(marca[\s:\.]*)',
            r'(?i)(seria\s*motor[\s:\.]*)',
            r'(?i)(seria\s*sasiu[\s:\.]*)',
            r'(?i)(cap\.?\s*cil[\s:\.]*)',
            r'(?i)(capacitate[\s:\.]*)',
            r'(?i)(tona[je]*[\s:\.]*)',
            r'(?i)(nr\.?\s*declaratie[\s:\.]*)',
            r'(?i)(nr\.?\s*matricol[\s:\.]*)',
            r'(?i)(nr\.?\s*inmatriculare[\s:\.]*)',
            r'(?i)(str\.?[\s:\.]*)',
            r'(?i)(bl\.?[\s:\.]*)',
            r'(?i)(sc\.?[\s:\.]*)',
            r'(?i)(et\.?[\s:\.]*)',
            r'(?i)(ap\.?[\s:\.]*)',
        ]
        seen = set()
        for pat in patterns:
            for m in re.finditer(pat, text):
                lbl = m.group(1).strip().rstrip(':').strip()
                if lbl and lbl.lower() not in seen:
                    seen.add(lbl.lower())
                    start = max(0, m.start()-30)
                    end = min(len(text), m.end()+60)
                    ctx = text[start:end].replace('\n',' ')
                    results.append((lbl, ctx))
        return results

    # ------------------------------------------------------------------
    # SEMNĂTURĂ
    # ------------------------------------------------------------------
    def _sig_draw(self):
        if not PIL_OK:
            return messagebox.showerror("Eroare", "Pillow nu este instalat!")
        win = tk.Toplevel(self)
        win.title("Pad Semnătură")
        win.geometry("500x220")
        win.resizable(False, False)
        win.transient(self)
        win.grab_set()

        ttk.Label(win, text="Semnați ținând click stânga și tragând:", font=('Segoe UI', 10, 'bold')).pack(pady=3)
        c = tk.Canvas(win, bg="white", width=480, height=130, relief=tk.SUNKEN, bd=2)
        c.pack(padx=10, pady=3)

        img = Image.new("RGBA", (480, 130), (255,255,255,0))
        draw = ImageDraw.Draw(img)
        last = [None, None]

        def paint(e):
            if last[0] is not None:
                c.create_line(last[0], last[1], e.x, e.y, fill="black", width=2, capstyle=tk.ROUND, smooth=True)
                draw.line([(last[0], last[1]), (e.x, e.y)], fill="black", width=2)
            last[0], last[1] = e.x, e.y
        def reset(e): last[0] = last[1] = None
        c.bind("<B1-Motion>", paint)
        c.bind("<ButtonRelease-1>", reset)

        def clear():
            c.delete("all")
            nonlocal img, draw
            img = Image.new("RGBA", (480,130), (255,255,255,0))
            draw = ImageDraw.Draw(img)
        def save():
            if img.getbbox() is None:
                return messagebox.showwarning("Atenție", "Semnătura e goală!")
            self.signature_img = img
            self.sig_lbl.config(text="Semnătură: Configurată ✓", foreground="green")
            self.status_var.set("Semnătură salvată")
            win.destroy()

        bf = ttk.Frame(win)
        bf.pack(pady=5)
        ttk.Button(bf, text="🗑️ Șterge", command=clear).pack(side=tk.LEFT, padx=5)
        ttk.Button(bf, text="💾 Salvează", command=save).pack(side=tk.LEFT, padx=5)

    def _sig_load_img(self):
        if not PIL_OK:
            return messagebox.showerror("Eroare", "Pillow nu este instalat!")
        p = filedialog.askopenfilename(filetypes=[("Imagini", "*.png *.jpg *.jpeg *.bmp *.gif")])
        if p:
            try:
                self.signature_img = Image.open(p).convert("RGBA")
                self.sig_lbl.config(text="Semnătură: Imagine încărcată ✓", foreground="green")
                self.status_var.set("Semnătură încărcată")
            except Exception as e:
                messagebox.showerror("Eroare", str(e))

    # ------------------------------------------------------------------
    # SALVARE
    # ------------------------------------------------------------------
    def _fill_save(self):
        if not self.current_file:
            return messagebox.showwarning("Atenție", "Încărcați un document!")

        data = []
        for f in self.manual_fields:
            lbl = f["label"].get().strip()
            val = f["value"].get().strip()
            if lbl and val:
                data.append({
                    "label": lbl, "value": val,
                    "x": f.get("x"), "y": f.get("y"),
                    "page": f.get("page", 0), "placed": f.get("placed", False)
                })

        has_sig = self.signature_img is not None
        if not data and not has_sig:
            return messagebox.showwarning("Atenție", "Nu există date sau semnătură!")

        out = filedialog.asksaveasfilename(
            title="Salvează documentul completat",
            defaultextension=".pdf" if self.current_doc_type=="pdf" else ".docx",
            filetypes=[("PDF", "*.pdf"), ("DOCX", "*.docx")]
        )
        if not out: return

        try:
            if self.current_doc_type == "pdf" and out.endswith(".pdf"):
                self._save_pdf_filled(out, data)
            elif self.current_doc_type == "docx" and out.endswith(".docx"):
                self._save_docx_filled(out, data)
            elif self.current_doc_type == "pdf" and out.endswith(".docx"):
                if not PDF2DOCX_OK:
                    return messagebox.showerror("Eroare", "Instalați pdf2docx!")
                tmp = out.replace(".docx", "_tmp.docx")
                cv = Converter(self.current_file)
                cv.convert(tmp, start=0, end=None)
                cv.close()
                self._save_docx_filled(out, data, template=tmp)
                os.remove(tmp)
            else:
                return messagebox.showerror("Eroare", "Combinație format nesuportată!")

            self.status_var.set(f"Salvat: {os.path.basename(out)}")
            messagebox.showinfo("Succes", f"Document salvat:\n{out}")
        except Exception as e:
            messagebox.showerror("Eroare salvare", str(e))
            self.status_var.set("Eroare la salvare")

    def _save_pdf_filled(self, out_path, data):
        if not PYMUPDF_OK:
            r = PdfReader(self.current_file)
            w = PdfWriter()
            for p in r.pages: w.add_page(p)
            meta = {"/Producer": "PyPDF Studio V2", "/ModDate": datetime.now().strftime("D:%Y%m%d%H%M%S")}
            if data:
                info = "; ".join([f"{d['label']}: {d['value']}" for d in data])[:300]
                meta["/Subject"] = info
            w.add_metadata(meta)
            with open(out_path, 'wb') as f: w.write(f)
            return

        doc = fitz.open(self.current_file)

        # Plasează fiecare câmp
        for d in data:
            if d.get("placed") and d.get("x") is not None:
                page = doc[d["page"]]
                x = d["x"] / self.zoom
                y = d["y"] / self.zoom
                page.insert_text((x, y), d["value"], fontsize=11, color=(0,0,0))

        # Plasează semnătura
        if self.signature_img and self.sig_x is not None:
            page = doc[self.sig_page] if self.sig_page < len(doc) else doc[-1]
            sig_tmp = out_path.replace(".pdf", "_sigtmp.png")
            self.signature_img.save(sig_tmp)
            x = self.sig_x / self.zoom
            y = self.sig_y / self.zoom
            rect = fitz.Rect(x-50, y-25, x+50, y+25)
            page.insert_image(rect, filename=sig_tmp)
            os.remove(sig_tmp)

        doc.save(out_path)
        doc.close()

    def _save_docx_filled(self, out_path, data, template=None):
        if not DOCX_OK:
            return messagebox.showerror("Eroare", "python-docx nu este instalat!")
        src = template if template else self.current_file
        doc = Document(src)

        for para in doc.paragraphs:
            full = para.text
            modified = False
            for d in data:
                if d["label"].lower() in full.lower():
                    para.clear()
                    run = para.add_run(f"{d['label']}: {d['value']}")
                    run.font.size = Pt(11)
                    modified = True
                    break
            if not modified and ('....' in full or '____' in full or '___' in full):
                for d in data:
                    para.clear()
                    para.add_run(f"{d['label']}: {d['value']}")
                    break

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for d in data:
                        if d["label"].lower() in cell.text.lower():
                            cell.text = f"{d['label']}: {d['value']}"

        if self.signature_img and PIL_OK:
            doc.add_paragraph()
            p = doc.add_paragraph("Semnătură: ")
            sig_tmp = out_path.replace(".docx", "_sigtmp.png")
            self.signature_img.save(sig_tmp)
            run = p.add_run()
            run.add_picture(sig_tmp, width=Inches(1.5))
            os.remove(sig_tmp)

        doc.save(out_path)

    # ================================================================
    # TAB 4: SECURITY
    # ================================================================
    def _tab_security(self):
        self.tab_sec = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_sec, text="  Securitate & Metadate  ")
        f = ttk.Frame(self.tab_sec, padding=15)
        f.pack(expand=True, fill="both")
        ttk.Label(f, text="Securitate PDF", style='Header.TLabel').pack(anchor=tk.W)

        pfrm = ttk.LabelFrame(f, text="Protejare cu parolă", padding=10)
        pfrm.pack(fill=tk.X, pady=8)
        self.sec_path = tk.StringVar()
        ttk.Entry(pfrm, textvariable=self.sec_path, state='readonly', width=55).pack(anchor=tk.W)
        ttk.Button(pfrm, text="📂 Selectează PDF", command=lambda: self._sec_load()).pack(anchor=tk.W, pady=3)
        ttk.Label(pfrm, text="Parolă utilizator:").pack(anchor=tk.W)
        self.sec_user = tk.StringVar()
        ttk.Entry(pfrm, textvariable=self.sec_user, show="*", width=40).pack(anchor=tk.W)
        ttk.Label(pfrm, text="Parolă proprietar (opțional):").pack(anchor=tk.W)
        self.sec_owner = tk.StringVar()
        ttk.Entry(pfrm, textvariable=self.sec_owner, show="*", width=40).pack(anchor=tk.W)
        ttk.Button(pfrm, text="🔒 Criptează", command=self._sec_encrypt).pack(anchor=tk.W, pady=8)

        mfrm = ttk.LabelFrame(f, text="Metadate", padding=10)
        mfrm.pack(fill=tk.BOTH, expand=True, pady=8)
        ttk.Button(mfrm, text="📂 Încarcă PDF", command=self._meta_load).pack(anchor=tk.W)
        self.meta_txt = scrolledtext.ScrolledText(mfrm, height=10, wrap=tk.WORD, font=('Consolas', 9))
        self.meta_txt.pack(fill=tk.BOTH, expand=True, pady=5)
        ttk.Button(mfrm, text="💾 Salvează metadate modificate", command=self._meta_save).pack(anchor=tk.W)

    def _sec_load(self):
        p = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if p: self.sec_path.set(p)

    def _sec_encrypt(self):
        p = self.sec_path.get()
        if not p: return messagebox.showwarning("Atenție", "Selectați PDF!")
        u = self.sec_user.get()
        if not u: return messagebox.showwarning("Atenție", "Introduceți parola!")
        out = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not out: return
        try:
            r = PdfReader(p)
            w = PdfWriter()
            for pg in r.pages: w.add_page(pg)
            if r.metadata: w.add_metadata(r.metadata)
            w.encrypt(user_password=u, owner_password=(self.sec_owner.get() or u))
            with open(out, 'wb') as f: w.write(f)
            self.status_var.set("PDF criptat")
            messagebox.showinfo("Succes", f"Protejat: {out}")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    def _meta_load(self):
        p = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if not p: return
        try:
            r = PdfReader(p)
            self.meta_file = p
            self.meta_txt.delete(1.0, tk.END)
            if r.metadata:
                for k,v in r.metadata.items():
                    self.meta_txt.insert(tk.END, f"{k}: {v}\n")
            else:
                self.meta_txt.insert(tk.END, "Fără metadate.")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    def _meta_save(self):
        if not getattr(self, 'meta_file', None):
            return messagebox.showwarning("Atenție", "Încărcați PDF!")
        out = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not out: return
        try:
            r = PdfReader(self.meta_file)
            w = PdfWriter()
            for pg in r.pages: w.add_page(pg)
            new_meta = {}
            for line in self.meta_txt.get(1.0, tk.END).strip().split("\n"):
                if ":" in line:
                    k,v = line.split(":", 1)
                    new_meta[k.strip()] = v.strip()
            if new_meta: w.add_metadata(new_meta)
            with open(out, 'wb') as f: w.write(f)
            self.status_var.set("Metadate actualizate")
            messagebox.showinfo("Succes", "Salvat!")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    # ================================================================
    # TAB 5: CONVERT
    # ================================================================
    def _tab_convert(self):
        self.tab_conv = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_conv, text="  Conversie  ")
        f = ttk.Frame(self.tab_conv, padding=15)
        f.pack(expand=True, fill="both")
        ttk.Label(f, text="Conversie Formate", style='Header.TLabel').pack(anchor=tk.W)

        a = ttk.LabelFrame(f, text="PDF → Word (DOCX)", padding=10)
        a.pack(fill=tk.X, pady=8)
        self.c_pdf = tk.StringVar()
        ttk.Entry(a, textvariable=self.c_pdf, state='readonly', width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(a, text="📂 PDF", command=lambda: self._c_sel("pdf")).pack(side=tk.LEFT, padx=3)
        ttk.Button(a, text="🔄 Convertește", command=self._c_pdf2docx).pack(side=tk.LEFT, padx=3)

        b = ttk.LabelFrame(f, text="Word (DOCX) → PDF", padding=10)
        b.pack(fill=tk.X, pady=8)
        self.c_docx = tk.StringVar()
        ttk.Entry(b, textvariable=self.c_docx, state='readonly', width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(b, text="📂 DOCX", command=lambda: self._c_sel("docx")).pack(side=tk.LEFT, padx=3)
        ttk.Button(b, text="🔄 Convertește", command=self._c_docx2pdf).pack(side=tk.LEFT, padx=3)
        ttk.Label(b, text="(necesită Microsoft Word instalat pe Windows)", foreground="gray").pack(side=tk.LEFT, padx=5)

        c = ttk.LabelFrame(f, text="OCR - Text din PDF scanat", padding=10)
        c.pack(fill=tk.BOTH, expand=True, pady=8)
        ttk.Label(c, text="Extrage text din PDF-uri scanate (necesită Tesseract):")
        self.ocr_txt = scrolledtext.ScrolledText(c, height=8, wrap=tk.WORD, font=('Consolas', 9))
        self.ocr_txt.pack(fill=tk.BOTH, expand=True, pady=5)
        ttk.Button(c, text="📂 Încarcă PDF scanat", command=self._c_ocr).pack(anchor=tk.W)

    def _c_sel(self, t):
        if t=="pdf":
            p = filedialog.askopenfilename(filetypes=[("PDF","*.pdf")])
            if p: self.c_pdf.set(p)
        else:
            p = filedialog.askopenfilename(filetypes=[("DOCX","*.docx")])
            if p: self.c_docx.set(p)

    def _c_pdf2docx(self):
        if not PDF2DOCX_OK:
            return messagebox.showerror("Eroare", "Instalați: pip install pdf2docx")
        p = self.c_pdf.get()
        if not p: return messagebox.showwarning("Atenție", "Selectați PDF!")
        out = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX","*.docx")])
        if not out: return
        try:
            self.status_var.set("Conversie în curs...")
            self.update()
            cv = Converter(p); cv.convert(out, start=0, end=None); cv.close()
            self.status_var.set("Conversie reusita")
            messagebox.showinfo("Succes", f"Salvat: {out}")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    def _c_docx2pdf(self):
        if not DOCX2PDF_OK:
            return messagebox.showerror("Eroare", "Instalați: pip install docx2pdf\nNecesită Microsoft Word instalat.")
        p = self.c_docx.get()
        if not p: return messagebox.showwarning("Atenție", "Selectați DOCX!")
        out = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF","*.pdf")])
        if not out: return
        try:
            self.status_var.set("Conversie în curs... (poate dura)")
            self.update()
            convert(p, out)
            self.status_var.set("Conversie reusita")
            messagebox.showinfo("Succes", f"Salvat: {out}")
        except Exception as e:
            messagebox.showerror("Eroare", f"Conversie eșuată:\n{str(e)}\n\nAsigurați-vă că Microsoft Word este instalat.")

    def _c_ocr(self):
        if not OCR_OK:
            return messagebox.showerror("Eroare", "Instalați: pip install pytesseract pdf2image")
        p = filedialog.askopenfilename(filetypes=[("PDF","*.pdf")])
        if not p: return
        try:
            self.status_var.set("OCR în curs...")
            self.update()
            imgs = convert_from_path(p)
            text = ""
            for i, img in enumerate(imgs):
                text += f"\n--- Pagina {i+1} ---\n"
                text += pytesseract.image_to_string(img, lang='ron+eng')
            self.ocr_txt.delete(1.0, tk.END)
            self.ocr_txt.insert(tk.END, text)
            self.status_var.set("OCR finalizat")
        except Exception as e: messagebox.showerror("Eroare", str(e))

    def _about(self):
        messagebox.showinfo("Despre",
            "PyPDF Studio V2\nProcesor Universal de Documente\n\n"
            "Funcții:\n"
            "• Unire & Divizare PDF\n"
            "• Completare formulare cu plasare pe pagină\n"
            "• Semnătură electronică gratuită\n"
            "• Zoom & Panning în previzualizare\n"
            "• Preview DOCX în mod text\n"
            "• Securitate & Metadate\n"
            "• Conversie PDF↔DOCX\n"
            "• OCR pentru scanări")


if __name__ == "__main__":
    app = PyPDFStudioV2()
    app.mainloop()
